import os
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from langchain_community.document_loaders import YoutubeLoader
import streamlit as st
from pathlib import Path

load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Logic functions
def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1", 
            file=audio_file
        )
    return transcription.text

def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": """
                You are a highly skilled AI trained in language comprehension and summarization. 
                I would like you to read the following text and summarize it into a concise abstract paragraph. 
                Aim to retain the most important points, providing a coherent and readable summary
                that could help a person understand the main points of the discussion without needing to read the entire text. 
                Please avoid unnecessary details or tangential points."""
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": """You are a proficient AI with a specialty in distilling information into key points.
                Based on the following text, identify and list the main points that were discussed or brought up. 
                These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion.
                Your goal is to provide a list that someone could read to quickly understand what was talked about."""
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": """
                You are an AI expert in analyzing conversations and extracting action items. 
                Please review the text and identify any tasks, assignments,
                or actions that were agreed upon or mentioned as needing to be done. 
                These could be tasks assigned to specific individuals, or general actions that the group has decided to take.
                Please list these action items clearly and concisely."""
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)

    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
    }

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()  # Add a line break between sections
    doc.save(filename)

def process_youtube_video(url):
    loader = YoutubeLoader.from_youtube_url(url, add_video_info=False)
    transcription = loader.load()[0].page_content

    minutes = meeting_minutes(transcription)

    audio_response = client.audio.speech.create(
        model="tts-1",
        voice="alloy",
        input=minutes['abstract_summary']
    )

    path_url = Path(url)
    audio_filename = f'audio_{path_url.stem}.mp3'
    audio_response.write_to_file(audio_filename)
    docx_filename = f'{path_url.stem}.docx'
    save_as_docx(minutes=minutes, filename=docx_filename)

    return transcription, minutes, audio_filename, docx_filename

# Streamlit UI
st.set_page_config(layout='wide')

def main():
    st.title("Meeting Minutes Generator from YouTube Video")
    path = st.text_input(label="Enter a YouTube URL")
    if path:
        try:
            transcription, minutes, audio_filename, docx_filename = process_youtube_video(path)
            
            st.video(path)

            with st.expander(label='Abstract Summary'):
                st.markdown(minutes.get("abstract_summary"))

            with st.expander('Key Points :key:'):
                st.write(minutes.get('key_points'))

            with st.expander('Action Items :🎯:'):
                st.write(minutes.get('action_items'))

            st.audio(audio_filename)
            c1,c2=st.columns([1,1.5])
            with c1:
                st.download_button(label='Download Audio Transcription', data=open(audio_filename, 'rb'), file_name=audio_filename)
            with c2:
                st.download_button(label='Download Overview', data=open(docx_filename, 'rb'), file_name=docx_filename)
        except Exception as e:
            st.error(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
